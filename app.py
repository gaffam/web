import json
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import feedparser
import language_tool_python
from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas

from PySide6.QtCore import QDate, QTimer, Qt
from PySide6.QtGui import QColor, QFont, QTextCharFormat, QTextCursor, QTextEdit
from PySide6.QtWidgets import (
    QApplication,
    QCalendarWidget,
    QFileDialog,
    QFrame,
    QHBoxLayout,
    QLabel,
    QListWidget,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QSplitter,
    QTextEdit,
    QVBoxLayout,
    QWidget,
    QLineEdit,
)

APP_DIR = Path.home() / ".velin_desk"
NOTES_DIR = APP_DIR / "notes"
APP_DIR.mkdir(exist_ok=True)
NOTES_DIR.mkdir(exist_ok=True)


@dataclass
class Note:
    note_id: str
    title: str
    body: str
    created_at: str
    updated_at: str


class VelinDesk(QMainWindow):
    """VELIN Desk / Scriptorium: kişisel premium yazı atölyesi."""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("VELIN Desk · Scriptorium")
        self.resize(1600, 980)

        self.current_note_id = None
        self.tool = language_tool_python.LanguageToolPublicAPI("tr-TR")
        self.last_snapshot = ""

        self._build_ui()
        self._apply_premium_theme()
        self._wire_timers()
        self.refresh_notes()

    def _build_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        root = QHBoxLayout(central)
        root.setContentsMargins(18, 18, 18, 18)
        root.setSpacing(14)

        splitter = QSplitter(Qt.Horizontal)
        root.addWidget(splitter)

        # Sol panel
        left = QFrame()
        left.setObjectName("leftPanel")
        left_l = QVBoxLayout(left)
        left_l.setContentsMargins(22, 20, 16, 20)

        left_title = QLabel("Le Carnet · Defter")
        left_title.setObjectName("panelTitle")
        self.note_list = QListWidget()
        self.note_list.itemClicked.connect(self.load_selected_note)

        self.new_btn = QPushButton("＋ Yeni Sayfa")
        self.new_btn.clicked.connect(self.new_note)

        left_l.addWidget(left_title)
        left_l.addWidget(self.note_list, 1)
        left_l.addWidget(self.new_btn)

        # Orta panel
        mid = QFrame()
        mid.setObjectName("midPanel")
        mid_l = QVBoxLayout(mid)
        mid_l.setContentsMargins(30, 28, 30, 28)
        mid_l.setSpacing(10)

        self.title_edit = QLineEdit()
        self.title_edit.setPlaceholderText("Başlık · Titre")
        self.title_edit.setObjectName("titleEdit")

        self.date_stamp = QLabel(datetime.now().strftime("%d %B %Y · %H:%M"))
        self.date_stamp.setObjectName("dateStamp")

        self.editor = QTextEdit()
        self.editor.setObjectName("mainEditor")
        self.editor.setPlaceholderText("Mürekkebin sayfayla buluştuğu o ilk cümleyi yaz…")
        self.editor.textChanged.connect(self.on_text_changed)

        mid_l.addWidget(self.title_edit)
        mid_l.addWidget(self.date_stamp)
        mid_l.addWidget(self.editor, 1)

        # Sağ panel
        right = QFrame()
        right.setObjectName("rightPanel")
        right_l = QVBoxLayout(right)
        right_l.setContentsMargins(18, 18, 18, 18)

        right_title = QLabel("L'Almanach")
        right_title.setObjectName("panelTitle")

        self.calendar = QCalendarWidget()
        self.calendar.setGridVisible(False)

        self.agenda = QListWidget()
        self.agenda.addItems([
            "09:00 · Sabah notları",
            "14:00 · Düzeltme turu",
            "20:30 · Günlük kapanış",
        ])

        self.open_btn = QPushButton("Dosya Aç (.md/.pdf/.docx)")
        self.open_btn.clicked.connect(self.open_file)
        self.export_md_btn = QPushButton("Export .md")
        self.export_md_btn.clicked.connect(self.export_md)
        self.export_pdf_btn = QPushButton("Export .pdf")
        self.export_pdf_btn.clicked.connect(self.export_pdf)
        self.export_docx_btn = QPushButton("Export .docx")
        self.export_docx_btn.clicked.connect(self.export_docx)

        right_l.addWidget(right_title)
        right_l.addWidget(self.calendar)
        right_l.addWidget(QLabel("Ajanda"))
        right_l.addWidget(self.agenda, 1)
        right_l.addWidget(self.open_btn)
        right_l.addWidget(self.export_md_btn)
        right_l.addWidget(self.export_pdf_btn)
        right_l.addWidget(self.export_docx_btn)

        splitter.addWidget(left)
        splitter.addWidget(mid)
        splitter.addWidget(right)
        splitter.setSizes([280, 980, 320])

    def _apply_premium_theme(self):
        self.setStyleSheet(
            """
            QMainWindow {
                background: qlineargradient(x1:0,y1:0,x2:1,y2:1,
                    stop:0 #f7f0e0, stop:0.45 #f3ead7, stop:1 #ecdfc8);
            }
            QFrame#leftPanel {
                background: #efe2cc;
                border: 1px solid #b08840;
                border-left: 16px solid #6a1f1f;
                border-radius: 10px;
            }
            QFrame#midPanel {
                background: #f3ead7;
                border: 1px solid #d3bc93;
                border-radius: 8px;
            }
            QFrame#rightPanel {
                background: #ede3d0;
                border: 1px solid #c5aa7e;
                border-radius: 10px;
            }
            QLabel#panelTitle {
                color: #2b1c0d;
                font-family: "Cormorant Garamond";
                font-size: 30px;
                font-weight: 700;
            }
            QLineEdit#titleEdit {
                font-family: "Cormorant Garamond";
                font-size: 42px;
                font-style: italic;
                color: #2b1c0d;
                border: none;
                border-bottom: 1px solid #b08840;
                background: transparent;
            }
            QLabel#dateStamp { color: #2a6a86; font-size: 14px; letter-spacing: 1px; }
            QTextEdit#mainEditor {
                background: #f3ead7;
                color: #2b1c0d;
                border: 1px solid #d8c39f;
                border-left: 3px solid #6a1f1f;
                border-radius: 6px;
                padding: 18px;
                font-family: "EB Garamond";
                font-size: 18px;
                line-height: 1.8;
                selection-background-color: #d8c39f;
            }
            QListWidget, QCalendarWidget {
                background: #f6eedf;
                border: 1px solid #c9af84;
                color: #2b1c0d;
            }
            QPushButton {
                background: #e1cda6;
                color: #2b1c0d;
                border: 1px solid #b08840;
                border-radius: 8px;
                padding: 8px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #2b1c0d;
                color: #f3ead7;
            }
            """
        )

    def _wire_timers(self):
        self.autosave_timer = QTimer(self)
        self.autosave_timer.timeout.connect(self.autosave)
        self.autosave_timer.start(5000)

        self.clock_timer = QTimer(self)
        self.clock_timer.timeout.connect(lambda: self.date_stamp.setText(datetime.now().strftime("%d %B %Y · %H:%M")))
        self.clock_timer.start(1000)

    def make_note(self) -> Note:
        now = datetime.utcnow().isoformat()
        note_id = self.current_note_id or datetime.utcnow().strftime("%Y%m%d%H%M%S")
        return Note(note_id, self.title_edit.text().strip() or "İsimsiz", self.editor.toPlainText(), now, now)

    def note_path(self, note_id: str) -> Path:
        return NOTES_DIR / f"{note_id}.json"

    def refresh_notes(self):
        self.note_list.clear()
        for p in sorted(NOTES_DIR.glob("*.json"), reverse=True):
            data = json.loads(p.read_text(encoding="utf-8"))
            self.note_list.addItem(f"{p.stem} · {data.get('title', 'İsimsiz')}")

    def new_note(self):
        self.current_note_id = None
        self.title_edit.clear()
        self.editor.clear()

    def load_selected_note(self, item):
        note_id = item.text().split(" · ", 1)[0]
        data = json.loads(self.note_path(note_id).read_text(encoding="utf-8"))
        self.current_note_id = note_id
        self.title_edit.setText(data.get("title", ""))
        self.editor.setPlainText(data.get("body", ""))

    def autosave(self):
        note = self.make_note()
        if not note.body.strip() and not note.title.strip():
            return
        payload = note.__dict__
        self.note_path(note.note_id).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        self.current_note_id = note.note_id
        self.statusBar().showMessage("Otomatik kaydedildi", 1200)
        self.refresh_notes()

    def on_text_changed(self):
        text = self.editor.toPlainText()
        if text == self.last_snapshot:
            return
        self.last_snapshot = text
        self.run_spellcheck_realtime(text)

    def run_spellcheck_realtime(self, text: str):
        matches = self.tool.check(text[:12000])
        selections = []
        for m in matches[:300]:
            sel = QTextEdit.ExtraSelection()
            fmt = QTextCharFormat()
            rule = (m.ruleIssueType or "").lower()
            if "typographical" in rule or "misspelling" in rule:
                fmt.setUnderlineColor(QColor("#c13c37"))
            elif "grammar" in rule:
                fmt.setUnderlineColor(QColor("#2a6a86"))
            else:
                fmt.setUnderlineColor(QColor("#b08840"))
            fmt.setUnderlineStyle(QTextCharFormat.WaveUnderline)
            sel.format = fmt
            c = self.editor.textCursor()
            c.setPosition(m.offset)
            c.setPosition(m.offset + m.errorLength, QTextCursor.KeepAnchor)
            sel.cursor = c
            selections.append(sel)
        self.editor.setExtraSelections(selections)

    def open_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Dosya Aç", "", "Documents (*.md *.pdf *.docx)")
        if not path:
            return
        p = Path(path)
        if p.suffix.lower() == ".md":
            self.editor.setPlainText(p.read_text(encoding="utf-8", errors="ignore"))
        elif p.suffix.lower() == ".pdf":
            rd = PdfReader(str(p))
            self.editor.setPlainText("\n".join(page.extract_text() or "" for page in rd.pages))
        elif p.suffix.lower() == ".docx":
            doc = Document(str(p))
            self.editor.setPlainText("\n".join(par.text for par in doc.paragraphs))

    def export_md(self):
        path, _ = QFileDialog.getSaveFileName(self, "Markdown Kaydet", "", "Markdown (*.md)")
        if not path:
            return
        content = f"# {self.title_edit.text().strip()}\n\n{self.editor.toPlainText()}"
        Path(path).write_text(content, encoding="utf-8")

    def export_docx(self):
        path, _ = QFileDialog.getSaveFileName(self, "Word Kaydet", "", "Word (*.docx)")
        if not path:
            return
        doc = Document()
        doc.add_heading(self.title_edit.text().strip(), level=1)
        doc.add_paragraph(self.editor.toPlainText())
        doc.save(path)

    def export_pdf(self):
        path, _ = QFileDialog.getSaveFileName(self, "PDF Kaydet", "", "PDF (*.pdf)")
        if not path:
            return
        c = canvas.Canvas(path, pagesize=A4)
        w, h = A4
        margin = 20 * mm

        def roman(n: int):
            vals = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"), (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
            out = ""
            for v, s in vals:
                while n >= v:
                    out += s
                    n -= v
            return out

        c.setFillColorRGB(0.95, 0.91, 0.84)
        c.rect(0, 0, w, h, fill=1, stroke=0)
        c.setFillColorRGB(0.17, 0.11, 0.05)
        c.setFont("Times-BoldItalic", 22)
        c.drawString(margin, h - margin, self.title_edit.text().strip() or "İsimsiz")

        text_obj = c.beginText(margin, h - margin - 18 * mm)
        text_obj.setFont("Times-Roman", 13)
        text_obj.setLeading(24)
        lines = self.editor.toPlainText().splitlines() or [""]
        page = 1
        for line in lines:
            if text_obj.getY() < margin + 25:
                c.drawRightString(w - margin, margin * 0.7, roman(page))
                c.drawText(text_obj)
                c.showPage()
                c.setFillColorRGB(0.95, 0.91, 0.84)
                c.rect(0, 0, w, h, fill=1, stroke=0)
                c.setFillColorRGB(0.17, 0.11, 0.05)
                text_obj = c.beginText(margin, h - margin)
                text_obj.setFont("Times-Roman", 13)
                text_obj.setLeading(24)
                page += 1
            text_obj.textLine(line)
        c.drawText(text_obj)
        c.drawRightString(w - margin, margin * 0.7, roman(page))
        c.save()


if __name__ == "__main__":
    app = QApplication([])
    app.setApplicationName("VELIN Desk")
    win = VelinDesk()
    win.show()
    app.exec()
